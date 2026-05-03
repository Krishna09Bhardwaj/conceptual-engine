import re
import httpx


def extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        import io
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)[:40000]
    elif ext in ("docx", "doc"):
        import io
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        # also grab tables
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
        return "\n".join(lines)[:40000]
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="replace")[:40000]
    else:
        raise ValueError(f"Unsupported file type: .{ext}")

_WHATSAPP_LINE = re.compile(
    r"^\[(\d{1,2}/\d{1,2}/\d{2,4}),\s*(\d{1,2}:\d{2}(?::\d{2})?)\]\s*([^:]+):\s*(.+)$"
)

_SYSTEM_PHRASES = [
    "Messages and calls are end-to-end encrypted",
    "changed the subject",
    "added ",
    "removed ",
    " left",
    "created group",
    "changed this group",
    "changed the group",
    "security code changed",
    "You deleted this message",
    "This message was deleted",
    "<Media omitted>",
]


def parse_whatsapp_txt(content: str, client_name: str) -> str:
    content = content[:50000]
    lines = content.splitlines()
    messages = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if any(phrase in stripped for phrase in _SYSTEM_PHRASES):
            continue
        m = _WHATSAPP_LINE.match(stripped)
        if m:
            date_str, time_str, sender, msg = m.groups()
            sender = sender.strip()[:60]
            msg = msg.strip()[:1000]
            messages.append(f"[{date_str} {time_str}] {sender}: {msg}")
    if not messages:
        return f"WhatsApp export for {client_name} (raw):\n{content[:8000]}"
    return f"WhatsApp conversation for {client_name}:\n" + "\n".join(messages)


async def fetch_fathom_transcript(url: str) -> str:
    import json, html as html_module
    url = url.strip()
    if not url.startswith("https://fathom.video/"):
        return "ERROR: URL must start with https://fathom.video/ — other URLs are not allowed."
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    }
    try:
        async with httpx.AsyncClient(timeout=20.0, follow_redirects=True) as client:
            resp = await client.get(url, headers=headers)
            resp.raise_for_status()
            page_html = resp.text

            # Fathom embeds page props as JSON in data-page attribute of #app div
            m = re.search(r'id="app"\s+data-page="([^"]+)"', page_html)
            if m:
                raw_json = html_module.unescape(m.group(1))
                try:
                    page_data = json.loads(raw_json)
                    props = page_data.get("props", {})
                    transcript_url = props.get("copyTranscriptUrl")
                    if transcript_url:
                        t_resp = await client.get(transcript_url, headers=headers)
                        t_resp.raise_for_status()
                        t_data = t_resp.json()
                        raw_html = t_data.get("html", "")
                        if raw_html:
                            # Strip HTML tags and clean up whitespace
                            clean = re.sub(r"<br\s*/?>", "\n", raw_html, flags=re.IGNORECASE)
                            clean = re.sub(r"<[^>]+>", "", clean)
                            clean = html_module.unescape(clean)
                            clean = re.sub(r"\n{3,}", "\n\n", clean).strip()
                            if len(clean) > 100:
                                return f"Fathom call transcript:\n{clean[:12000]}"
                except (json.JSONDecodeError, KeyError):
                    pass

            return "Could not extract transcript. The link may be private or the recording is still processing."

    except httpx.HTTPStatusError as e:
        return f"ERROR fetching Fathom link: HTTP {e.response.status_code}"
    except httpx.TimeoutException:
        return "ERROR: Request timed out fetching Fathom link."
    except Exception as e:
        return f"ERROR fetching Fathom link: {str(e)[:200]}"
